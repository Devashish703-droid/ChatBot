# ---------------------------
# Imports & helper functions
# ---------------------------
import difflib
import re
from docx import Document
from sentence_transformers import SentenceTransformer
import numpy as np

# ---------------------------
# Import your QA data (converted to .py)
# ---------------------------
from qa_data import qa
# from web_conect import <functions_if_needed>  # import any functions you need

# ---------------------------
# Word doc reader
# ---------------------------
def load_docx_knowledge(file_path):
    doc = Document(file_path)
    knowledge = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            knowledge.append((text, style))
    return knowledge

def clean_text(text):
    text = text.lower()
    text = re.sub(r"[^a-z\s]", "", text)
    return text.strip()

def search_doc_knowledge(query, knowledge):
    query_clean = clean_text(query)
    headings = [(i, k[0]) for i, k in enumerate(knowledge) if "Heading" in k[1]]
    cleaned_headings = [clean_text(h[1]) for h in headings]

    best_match = difflib.get_close_matches(query_clean, cleaned_headings, n=1, cutoff=0.5)
    if not best_match:
        for ch in cleaned_headings:
            if query_clean in ch:
                best_match = [ch]
                break
    if not best_match:
        return None

    for idx, heading_text in headings:
        if clean_text(heading_text) == best_match[0]:
            section = []
            for j in range(idx, len(knowledge)):
                next_text, next_style = knowledge[j]
                if j != idx and "Heading" in next_style and next_text.strip():
                    break
                section.append(next_text)
            if len(section) <= 1:
                return None
            return "\n".join(section)
    return None

# ---------------------------
# Q&A helper
# ---------------------------
def get_best_match(user_input, questions):
    matches = difflib.get_close_matches(user_input.lower(), questions, n=1, cutoff=0.6)
    return matches[0] if matches else None

# ---------------------------
# Semantic Section Reader
# ---------------------------
def read_sections(file_path):
    doc = Document(file_path)
    sections = {}
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = (
            para.style.name.startswith('Heading') or
            (para.runs and para.runs[0].bold) or
            text.isupper() or
            text.endswith(':')
        )
        if is_heading:
            if current_heading and current_content:
                sections[current_heading] = " ".join(current_content).strip()
            current_heading = text
            current_content = []
        else:
            if current_heading is None:
                current_heading = text
            else:
                current_content.append(text)
    if current_heading and current_content:
        sections[current_heading] = " ".join(current_content).strip()
    if not sections:
        sections["Document"] = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    return sections

# ---------------------------
# Load Knowledge Base & Embeddings
# ---------------------------
knowledge = load_docx_knowledge("KasturiAssist Product Documentation.docx")
sections = read_sections("KasturiAssist Product Documentation.docx")
headings = list(sections.keys())
contents = list(sections.values())
model = SentenceTransformer('all-MiniLM-L6-v2')
heading_embeddings = np.array(model.encode(headings))

# ------------------------
# Document Summary
# ------------------------
def summarize_document():
    summary_sentences = []
    for content in contents:
        sentences = content.split('. ')
        summary_sentences.append('. '.join(sentences[:2]).strip())
    return "\n".join(summary_sentences)

# ------------------------
# Document Index / Tree
# ------------------------
def index_document():
    tree = []
    for heading in headings:
        parts = [p.strip() for p in heading.replace('-', ':').split(':') if p.strip()]
        tree.append(" > ".join(parts))
    return "\n".join(tree)

# ------------------------
# Chatbot answer function
# ------------------------
def get_answer(user_query, top_n=1):
    user_query_lower = user_query.lower()

    # Special commands
    if "summarize" in user_query_lower or "explain whole document" in user_query_lower:
        return summarize_document()
    if "index" in user_query_lower or "topic tree" in user_query_lower:
        return index_document()

    # 1️⃣ Check Q&A dictionary first
    if 'qa' in globals():
        best_match = get_best_match(user_query, qa.keys())
        if best_match:
            return qa[best_match]

    # 2️⃣ Check Word doc fuzzy heading search
    doc_answer = search_doc_knowledge(user_query, knowledge)
    if doc_answer:
        lines = doc_answer.split("\n")
        if len(lines) > 6:
            lines = lines[:6] + ["• ...more details in document"]
        return "\n".join(lines)

    # 3️⃣ Semantic search fallback
    query_emb = np.array(model.encode([user_query]))[0]
    similarities = np.dot(heading_embeddings, query_emb) / (
        np.linalg.norm(heading_embeddings, axis=1) * np.linalg.norm(query_emb)
    )
    top_indices = np.argsort(similarities)[::-1][:top_n]
    if len(top_indices) == 0 or similarities[top_indices[0]] < 0.4:
        return "Sorry, I could not find relevant information in the document."
    answers = [contents[i] for i in top_indices]
    return "\n\n".join(answers)
